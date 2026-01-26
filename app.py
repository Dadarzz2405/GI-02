from utils import can_mark_attendance, is_core_user
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, abort, Response, Blueprint
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from flask_bcrypt import Bcrypt
import os
from models import Pic, db, User, Session, Attendance, Notulensi
from datetime import datetime, date, timezone, timedelta
from ummalqura.hijri_date import HijriDate
import json
from werkzeug.utils import secure_filename
from ai import call_chatbot_groq
from flask_migrate import Migrate
import csv
from io import TextIOWrapper, StringIO, BytesIO
from docx import Document
from formatter import format_attendance
from summarizer import summarize_notulensi

UPLOAD_FOLDER = 'static/uploads/profiles'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'webp'}

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY')
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL')
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
db.init_app(app)
bcrypt = Bcrypt(app)

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
migrate = Migrate(app, db)

attendance_bp = Blueprint("attendance", __name__)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email'].strip().lower()
        password = request.form['password']
        user = User.query.filter_by(email=email).first()
        if user and bcrypt.check_password_hash(user.password, password):
            login_user(user)
            if user.must_change_password:
                return redirect(url_for('profile'))
            else:
                if user.role in ['admin', 'ketua', 'pembina']:
                    return redirect(url_for('dashboard_admin'))
                else:
                    return redirect(url_for('dashboard_member'))
        else:
            flash('Invalid email or password', 'error')
        
    return render_template('login.html')


@app.route('/')
def home():
    if current_user.is_authenticated:
        if current_user.role in ['admin', 'ketua', 'pembina']:
            return redirect(url_for('dashboard_admin'))
        else:
            return redirect(url_for('dashboard_member'))
    else:
        return redirect(url_for('login'))

@app.route('/dashboard_admin')
@login_required
def dashboard_admin():
    if not current_user.role in ['admin', 'ketua', 'pembina']:
        return "Access denied"
    return render_template('dashboard_admin.html')

@app.route('/dashboard_member')
@login_required
def dashboard_member():
    if current_user.role in ['admin', 'ketua', 'pembina']:
        return redirect(url_for('dashboard_admin'))
    return render_template('dashboard_member.html')

@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        existing_user = User.query.filter_by(username=username).first()
        if existing_user and existing_user.id != current_user.id:
            flash('Username already taken', 'error')
            return redirect(url_for('profile')) 

        current_user.username = username
        
        if password: 
            current_user.password = bcrypt.generate_password_hash(password).decode('utf-8')
        
        db.session.commit()
        flash('Profile updated successfully', 'success')
        return redirect(url_for('profile'))
    
    return render_template('profile.html')

@app.route('/member-list')
@login_required
def member_list():
    users = User.query.all()
    return render_template('member_list.html', users=users)

@app.route('/create-session', methods=['GET', 'POST'])
@login_required
def create_session():
    if current_user.role not in ['admin', 'ketua', 'pembina']:
        return "Access denied"
    if request.method == 'POST':
        name = request.form['name']
        date = request.form['date']

        new_session = Session(name=name, date=date)
        db.session.add(new_session)
        db.session.commit()
        return redirect(url_for('dashboard_admin'))
    pics = Pic.query.all()
    return render_template('create_session.html', pics=pics)

from datetime import date
@app.route('/api/attendance', methods=['POST'])
@login_required
def api_attendance():
    data = request.get_json()

    session_id = data.get("session_id")
    user_id = data.get("user_id")
    status = data.get("status")

    if not all([session_id, user_id, status]):
        return jsonify({"error": "invalid_data"}), 400

    session = Session.query.get_or_404(session_id)

    if session.is_locked:
        return jsonify({"error": "session_locked"}), 403

    if not can_mark_attendance(current_user, session.pic_id):
        return jsonify({"error": "forbidden"}), 403


    record = Attendance.query.filter_by(
        session_id=session_id,
        user_id=user_id,
        attendance_type='regular'
    ).first()

    if record:
        return jsonify({"error": "already_marked"}), 409

    wib = timezone(timedelta(hours=7))
    attendance = Attendance(
        session_id=session_id,
        user_id=user_id,
        status=status,
        attendance_type='regular',
        timestamp=datetime.now(wib)
    )

    db.session.add(attendance)
    db.session.commit()

    return jsonify({"success": True})

@app.route("/attendance/core")
@login_required
def attendance_core():
    if not is_core_user(current_user):
        abort(403)

    sessions = Session.query.order_by(Session.date.desc()).all()
    users = User.query.filter(User.role.in_(["admin", "ketua"])).all()
    core_users = users

    return render_template(
        "attendance_mark_core.html",
        sessions=sessions,
        users=users,
        core_users=core_users
    )

@app.route("/api/attendance/core", methods=["POST"])
@login_required
def api_attendance_core():
    if not is_core_user(current_user):
        return jsonify({"error": "forbidden"}), 403

    data = request.get_json()
    session_id = data.get("session_id")
    user_id = data.get("user_id")
    status = data.get("status")

    if not all([session_id, user_id, status]):
        return jsonify({"error": "invalid_data"}), 400

    session = Session.query.get_or_404(session_id)

    if session.is_locked:
        return jsonify({"error": "session_locked"}), 403

    user = User.query.get_or_404(user_id)

    if not is_core_user(user):
        return jsonify({"error": "not_core_user"}), 400

    exists = Attendance.query.filter_by(
        session_id=session_id,
        user_id=user_id,
        attendance_type="core"
    ).first()

    if exists:
        return jsonify({"error": "already_marked"}), 409

    wib = timezone(timedelta(hours=7))
    att = Attendance(
        session_id=session_id,
        user_id=user_id,
        status=status,
        attendance_type="core",
        timestamp=datetime.now(wib)
    )

    db.session.add(att)
    db.session.commit()

    return jsonify({"success": True})

@app.route("/api/session/<int:session_id>/status", methods=["GET"])
@login_required
def get_session_status(session_id):
    """Get session lock status"""
    session = Session.query.get_or_404(session_id)
    return jsonify({
        "is_locked": session.is_locked,
        "session_id": session.id,
        "name": session.name
    })

@app.route("/api/session/<int:session_id>/lock", methods=["POST"])
@login_required
def lock_session(session_id):
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"error": "forbidden"}), 403

    session = Session.query.get_or_404(session_id)
    session.is_locked = True
    db.session.commit()

    return jsonify({"locked": True})

@app.route('/pics', methods=['GET', 'POST'])
@login_required
def manage_pics():
    if current_user.role not in ['admin', 'ketua', 'pembina']:
        abort(403)

    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        if not name:
            flash("PIC name cannot be empty", "error")
            return redirect(url_for('manage_pics'))

        existing = Pic.query.filter_by(name=name).first()
        if existing:
            flash("PIC already exists!", "error")
        else:
            new_pic = Pic(name=name)
            db.session.add(new_pic)
            db.session.commit()
            flash(f"PIC '{name}' created!", "success")

        return redirect(url_for('manage_pics'))

    pics = Pic.query.all()
    return render_template('manage_pic.html', pics=pics)

@app.route("/export/attendance/<int:session_id>")
@login_required
def export_attendance_csv(session_id):
    if current_user.role not in ["admin", "ketua", "pembina"]:
        abort(403)

    session = Session.query.get_or_404(session_id)

    records = (
        db.session.query(
            Attendance,
            User.name,
            User.email,
            User.role
        )
        .join(User, Attendance.user_id == User.id)
        .filter(Attendance.session_id == session_id)
        .order_by(User.name)
        .all()
    )

    if not records:
        flash("No attendance records found for this session", "warning")
        return redirect(url_for('attendance_mark'))

    wib = timezone(timedelta(hours=7))
    
    doc = Document()
    doc.add_heading(f'Attendance Report: {session.name}', 0)
    doc.add_paragraph(f'Date: {session.date}')
    doc.add_paragraph(f'Total Attendees: {len(records)}')
    doc.add_paragraph('')
    
    summary = {
        'present': sum(1 for a, _, _, _ in records if a.status == 'present'),
        'absent': sum(1 for a, _, _, _ in records if a.status == 'absent'),
        'excused': sum(1 for a, _, _, _ in records if a.status == 'excused'),
        'late': sum(1 for a, _, _, _ in records if a.status == 'late'),
    }
    
    doc.add_heading('Summary', level=1)
    summary_table = doc.add_table(rows=5, cols=2)
    summary_table.style = 'Light Grid Accent 1'
    
    summary_data = [
        ('Status', 'Count'),
        ('Present', str(summary['present'])),
        ('Absent', str(summary['absent'])),
        ('Excused', str(summary['excused'])),
        ('Late', str(summary['late']))
    ]
    
    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value
    
    doc.add_paragraph('')
    
    doc.add_heading('Detailed Records', level=1)
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Name'
    header_cells[1].text = 'Role'
    header_cells[2].text = 'Status'
    header_cells[3].text = 'Time'
    header_cells[4].text = 'Type'
    
    for attendance, name, email, role in records:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = role.capitalize()
        row_cells[2].text = attendance.status.capitalize()
        formatted_time = attendance.timestamp.astimezone(wib).strftime("%H:%M")
        row_cells[3].text = formatted_time
        row_cells[4].text = attendance.attendance_type.capitalize()
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    filename = f"attendance_{session.name.replace(' ', '_')}_{session.date}.docx"

    return Response(
        bio,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        }
    )

@app.route('/pic/delete/<int:id>', methods=['POST'])
@login_required
def delete_pic(id):
    if current_user.role != 'admin':
        abort(403)

    pic = Pic.query.get_or_404(id)

    for user in pic.members:
        user.pic_id = None
        user.can_mark_attendance = False  

    db.session.delete(pic)
    db.session.commit()

    flash("PIC deleted and permissions revoked", "success")
    return redirect(url_for('manage_pics'))

@app.route('/attendance-history')
@login_required
def attendance_history():
    records = Attendance.query.filter_by(user_id=current_user.id).all()

    summary = {
        'present': sum(1 for r in records if r.status=='present'),
        'absent': sum(1 for r in records if r.status=='absent'),
        'excused': sum(1 for r in records if r.status=='excused')
    }

    return render_template('attendance_history.html', records=records, summary=summary)

@app.route('/attendance-history-admin')
@login_required
def attendance_history_admin():
    if current_user.role not in ['admin', 'ketua', 'pembina']:
        return redirect(url_for('invalid_credential')) 
    users = User.query.filter(User.role=='member').all()
    return render_template('attendance_history_admin.html', users=users)

@app.route('/attendance-history-admin/<int:user_id>')
@login_required
def attendance_history_admin_view(user_id):
    if current_user.role not in ['admin', 'ketua', 'pembina']:
        return redirect(url_for('invalid_credential'))
    
    selected_user = User.query.get_or_404(user_id)

    records = Attendance.query.filter_by(user_id=user_id).all()
    
    summary = {
        'present': sum(1 for r in records if r.status=='present'),
        'absent': sum(1 for r in records if r.status=='absent'),
        'excused': sum(1 for r in records if r.status=='excused')
    }
    return render_template('attendance_history_admin_view.html', user=selected_user, records=records, summary=summary)

@app.route('/attendance-mark')
@login_required
def attendance_mark():
    if current_user.role not in ['admin', 'ketua']:
        abort(403)
    sessions = Session.query.all()
    users = User.query.filter(User.role == 'member').all()
    return render_template('attendance_mark_core.html', sessions=sessions, users=users)

@app.route('/attendance', methods=['GET', 'POST'])
@login_required
def attendance():
    if current_user.role in ['admin', 'ketua', 'pembina']:
        return redirect(url_for('attendance_mark'))
    
    is_pic = current_user.can_mark_attendance
    if not is_pic and current_user.role != 'member':
        abort(403)
    
    if request.method == 'POST':
        session_id = request.form.get('session_id')
        if not session_id:
            flash('Session not selected', 'error')
            return redirect(url_for('attendance'))
        session = Session.query.get_or_404(session_id)
        if session.is_locked:
            flash('Session is locked', 'error')
            return redirect(url_for('attendance'))
        
        if is_pic:
            if not can_mark_attendance(current_user, session.pic_id):
                abort(403)
            members = User.query.filter_by(pic_id=current_user.id).all()
        else:
            members = User.query.filter(User.role == 'member').all()
        
        wib = timezone(timedelta(hours=7))
        for member in members:
            status = request.form.get(f'status_{member.id}')
            if status:
                existing = Attendance.query.filter_by(session_id=session_id, user_id=member.id, attendance_type='regular').first()
                if not existing:
                    attendance = Attendance(session_id=session_id, user_id=member.id, status=status, attendance_type='regular', timestamp=datetime.now(wib))
                    db.session.add(attendance)
        db.session.commit()
        flash('Attendance saved', 'success')
        return redirect(url_for('attendance'))
    
    selected_session_id = request.args.get('session_id')
    
    if is_pic:
        sessions = Session.query.filter_by(pic_id=current_user.id).all()
        members = User.query.filter_by(pic_id=current_user.id).all()
    else:
        sessions = Session.query.all()
        members = User.query.filter(User.role == 'member').all()
    
    attendance_map = {}
    if selected_session_id:
        attendances = Attendance.query.filter_by(session_id=selected_session_id, attendance_type='regular').all()
        for a in attendances:
            attendance_map[a.user_id] = a.status
    return render_template('attendance.html', sessions=sessions, members=members, selected_session_id=selected_session_id, attendance_map=attendance_map)

@app.route('/logout')
@login_required
def logout(): 
    logout_user()
    return redirect(url_for('login'))

@app.route('/change-password', methods=['GET', 'POST'])
@login_required
def change_password():
    if request.method == 'POST':
        old_password = request.form['old_password']
        new_password = request.form['new_password']
        confirm_password = request.form['confirm_password']

        if not bcrypt.check_password_hash(current_user.password, old_password):
            flash("Incorrect current password.", "danger")
        elif new_password != confirm_password:
            flash("New passwords don't match.", "danger")
        else:
            current_user.password = bcrypt.generate_password_hash(new_password).decode('utf-8')
            current_user.must_change_password = False
            db.session.commit()
            flash("Password updated successfully!", "success")
            if current_user.role in ['admin', 'ketua', 'pembina']:
                return redirect(url_for('dashboard_admin'))
            else:
                return redirect(url_for('dashboard_member'))
    return render_template('change_password.html')

@app.route("/profile/upload_pfp", methods=['POST'])
@login_required
def upload_pfp():
    file = request.files.get('pfp')

    if not file or file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('profile'))
    
    filename_attr = getattr(file, 'filename', '')
    if not filename_attr or not allowed_file(filename_attr):
        flash('Invalid file type. Allowed types: png, jpg, jpeg, webp', 'error')
        return redirect(url_for('profile'))
    
    # Check file size (5MB limit)
    file.seek(0, os.SEEK_END)
    file_size = file.tell()
    file.seek(0)
    
    if file_size > 5 * 1024 * 1024:  # 5MB
        flash('File size too large. Maximum 5MB allowed.', 'error')
        return redirect(url_for('profile'))
    
    # Read file as binary
    image_data = file.read()
    
    # Store in database
    current_user.profile_picture_data = image_data
    current_user.profile_picture_filename = secure_filename(filename_attr)
    
    db.session.commit()
    flash('Profile picture updated successfully', 'success')
    return redirect(url_for('profile'))

# Add route to serve profile pictures from database:
@app.route("/profile-picture/<int:user_id>")
def serve_profile_picture(user_id):
    user = User.query.get_or_404(user_id)
    
    if user.profile_picture_data:
        # Determine mime type from filename
        filename = user.profile_picture_filename or 'image.png'
        ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else 'png'
        
        mime_types = {
            'png': 'image/png',
            'jpg': 'image/jpeg',
            'jpeg': 'image/jpeg',
            'webp': 'image/webp'
        }
        
        mime_type = mime_types.get(ext, 'image/png')
        
        return Response(user.profile_picture_data, mimetype=mime_type)
    else:
        # Serve default image
        default_path = os.path.join('static', 'uploads', 'profiles', 'default.png')
        if os.path.exists(default_path):
            with open(default_path, 'rb') as f:
                return Response(f.read(), mimetype='image/png')
        else:
            abort(404)

ISLAMIC_HOLIDAYS = {
    # Muharram
    "01-01": "Islamic New Year",
    "01-09": "Day of Tasua",
    "01-10": "Day of Ashura",

    # Rabi' al-Awwal
    "03-12": "Mawlid al-Nabi",

    # Rajab
    "07-01": "Start of Rajab",
    "07-27": "Isra and Mi'raj",

    # Sha'ban
    "08-15": "Mid-Sha'ban (Laylat al-Bara'ah)",

    # Ramadan
    "09-01": "Start of Ramadan",
    "09-17": "Nuzul al-Qur'an",
    "09-21": "Laylat al-Qadr (possible)",
    "09-23": "Laylat al-Qadr (possible)",
    "09-25": "Laylat al-Qadr (possible)",
    "09-27": "Laylat al-Qadr (possible)",
    "09-29": "Laylat al-Qadr (possible)",

    # Shawwal
    "10-01": "Eid al-Fitr",
    "10-02": "Eid al-Fitr ",

    # Dhu al-Qi'dah
    "11-01": "Start of Dhuqa'dah",

    # Dhu al-Hijjah
    "12-01": "Start of Dhu al-Hijjah",
    "12-08": "Day of Tarwiyah",
    "12-09": "Day of Arafah",
    "12-10": "Eid al-Adha",
    "12-11": "Days of Tashreeq",
    "12-12": "Days of Tashreeq",
    "12-13": "Days of Tashreeq",
}

@app.route("/calendar")
@login_required
def calendar():
    
    return render_template("calendar.html")

def get_hijri_date(gregorian_date):
    try:
        g = datetime.strptime(gregorian_date, "%Y-%m-%d").date()
        h = HijriDate(g.year, g.month, g.day, gr=True)
        return f"{h.day} {h.month_name} {h.year} H"
    except Exception:
        return ""

def get_hijri_key_from_gregorian(g_date: date):
    h = HijriDate(g_date.year, g_date.month, g_date.day, gr=True)
    return f"{h.month:02d}-{h.day:02d}", h


    
@app.route('/api/dashboard_calendar')
@login_required
def api_dashboard_calendar():
    sessions = Session.query.all()
    calendar_events = []

    for session in sessions:
        hijri_date = get_hijri_date(session.date)
        calendar_events.append({
            'title': f"{session.name} ({hijri_date})",
            'start': session.date,
            'extendedProps': {
                'type': 'rohis_session'
            }
        })

    today = date.today()
    start_year = today.year - 1
    end_year = today.year + 1

    current = date(start_year, 1, 1)
    end = date(end_year, 12, 31)

    while current <= end:
        hijri_key, hijri = get_hijri_key_from_gregorian(current)

        if hijri_key in ISLAMIC_HOLIDAYS:
            calendar_events.append({
                'title': f"{ISLAMIC_HOLIDAYS[hijri_key]} ({hijri.day} {hijri.month_name} {hijri.year} H)",
                'start': current.isoformat(),
                'allDay': True,
                'backgroundColor': '#1e88e5',
                'borderColor': '#1565c0',
                'textColor': '#ffffff',
                'extendedProps': {
                    'type': 'islamic_holiday',
                    'hijri': f"{hijri.day} {hijri.month_name} {hijri.year} H"
                }
            })

        current = current.fromordinal(current.toordinal() + 1)
    return jsonify(calendar_events)

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
@app.route("/chat", methods=["POST"])
@login_required
def chat():
    data = request.get_json()
    user_message = data.get("message", "").strip()

    if not user_message:
        return jsonify({"reply": "Please type a question."})

    try:
        reply = call_chatbot_groq(user_message)
    except Exception as e:
        print("CHATBOT ERROR:", e)
        reply = "Error occurred. Check server logs."


    return jsonify({"reply": reply})

@app.route('/pic-management', methods=['GET', 'POST'])
@login_required
def pic_management():

    if current_user.role not in ['admin', 'ketua', 'pembina']:
        abort(403)

    pics = Pic.query.all()
    users = User.query.filter_by(role='member').all()

    if request.method == 'POST':
        user_ids = request.form.getlist('user_ids')
        pic_id = request.form.get('pic_id')
        marker_id = request.form.get('marker_id')

        if not user_ids or not pic_id:
            flash("Please select users and a PIC.", "danger")
            return redirect(url_for('pic_management'))

        try:
            user_ids = [int(uid) for uid in user_ids]
            pic_id = int(pic_id)
        except (ValueError, TypeError):
            flash("Invalid input values.", "danger")
            return redirect(url_for('pic_management'))

        try:
            marker_id = int(marker_id) if marker_id else None
        except (ValueError, TypeError):
            marker_id = None

        User.query.filter_by(pic_id=pic_id, can_mark_attendance=True).update({"can_mark_attendance": False})

        for uid in user_ids:
            user = User.query.get(uid)
            if user:
                user.pic_id = pic_id
                user.can_mark_attendance = True if (marker_id and uid == marker_id) else False
                db.session.add(user)

        db.session.commit()

        flash(f"Assigned {len(user_ids)} members to the PIC.", "success")
        return redirect(url_for('pic_management'))

    return render_template(
        'pic_management.html',
        pics=pics,
        users=users
    )

@app.route("/api/notulensi/<int:session_id>", methods=["POST"])
@login_required
def save_notulensi(session_id):
    if current_user.role not in ['admin', 'ketua', 'pembina']:
        abort(403)
    data = request.get_json()
    content = data.get("content", "").strip()

    if not content or content in ['<p><br></p>', '<p></p>', '<div><br></div>', '<div></div>']:
        return jsonify({"error": "Content cannot be empty"}), 400

    note = Notulensi.query.filter_by(session_id=session_id).first()

    if note:
        note.content = content
        note.updated_at = datetime.utcnow()
    else:
        note = Notulensi(session_id=session_id, content=content)
        db.session.add(note)

    db.session.commit()
    return jsonify({"success": True})

@app.route("/api/notulensi/<int:notulensi_id>", methods=["DELETE"])
@login_required
def delete_notulensi(notulensi_id):
    if current_user.role not in ['admin', 'ketua', 'pembina']:
        return jsonify({"error": "forbidden"}), 403
    
    note = Notulensi.query.get_or_404(notulensi_id)
    db.session.delete(note)
    db.session.commit()
    
    return jsonify({"success": True})

@app.route("/notulensi-list")
@login_required
def notulensi_list():
    sessions = Session.query.order_by(Session.date.desc()).all()
    notulensis = Notulensi.query.all()
    notulensi_dict = {n.session_id: n for n in notulensis}
    
    return render_template("notulensi_list.html", sessions=sessions, notulensi_dict=notulensi_dict)

@app.route("/notulensi/<int:session_id>")
@login_required
def notulensi(session_id):
    session = Session.query.get_or_404(session_id)
    note = Notulensi.query.filter_by(session_id=session_id).first()
    
    can_edit = current_user.role in ['admin', 'ketua', 'pembina']
    
    return render_template("notulensi.html", session=session, note=note, can_edit=can_edit)

@app.route("/notulensi/view/<int:notulensi_id>")
@login_required
def notulensi_view(notulensi_id):
    """View full notulensi content (read-only for members)"""
    note = Notulensi.query.get_or_404(notulensi_id)
    session = Session.query.get_or_404(note.session_id)
    
    can_edit = current_user.role in ['admin', 'ketua', 'pembina']
    
    return render_template("notulensi_view.html", session=session, note=note, can_edit=can_edit)

@app.route('/api/news-feed')
@login_required
def news_feed():
    """
    News feed API with comprehensive error handling and fallbacks
    """
    try:
        today = datetime.now().date()
        
        # Get upcoming sessions
        upcoming_sessions = Session.query.filter(
            Session.date >= str(today)
        ).order_by(Session.date.asc()).limit(3).all()
        
        # Get recent notulensi
        recent_notulensi = (
            db.session.query(Notulensi, Session)
            .join(Session, Notulensi.session_id == Session.id)
            .order_by(Notulensi.updated_at.desc())
            .limit(3)
            .all()
        )
        
        # Process upcoming sessions
        upcoming_data = []
        for session in upcoming_sessions:
            try:
                upcoming_data.append({
                    'id': session.id,
                    'name': session.name,
                    'date': session.date,
                    'pic': session.pic.name if session.pic else 'No PIC assigned'
                })
            except Exception as e:
                print(f"Error processing session {session.id}: {e}")
                continue
        
        # Process recent notulensi with better error handling
        recent_data = []
        for notulensi, session in recent_notulensi:
            try:
                # Try to generate summary, but have multiple fallbacks
                summary = "Meeting notes available."
                
                if notulensi and notulensi.content:
                    try:
                        # Only try to summarize if GROQ_API_KEY exists
                        if os.environ.get("GROQ_API_KEY"):
                            summary = summarize_notulensi(notulensi.content)
                        else:
                            # Fallback: Use first 150 characters of cleaned text
                            from html import unescape
                            import re
                            clean_text = re.sub('<[^<]+?>', '', notulensi.content)
                            clean_text = unescape(clean_text).strip()
                            if len(clean_text) > 150:
                                summary = clean_text[:150] + "..."
                            else:
                                summary = clean_text if clean_text else "Meeting notes available."
                    except Exception as sum_error:
                        print(f"Summarization error for notulensi {notulensi.id}: {sum_error}")
                        # Fallback to manual text extraction
                        try:
                            from html import unescape
                            import re
                            clean_text = re.sub('<[^<]+?>', '', notulensi.content)
                            clean_text = unescape(clean_text).strip()
                            summary = clean_text[:150] + "..." if len(clean_text) > 150 else (clean_text or "Meeting notes available.")
                        except:
                            summary = "Meeting notes available."
                
                recent_data.append({
                    'id': notulensi.id,
                    'session_name': session.name,
                    'session_date': session.date,
                    'summary': summary,
                    'updated_at': notulensi.updated_at.strftime('%d %b %Y') if notulensi.updated_at else notulensi.created_at.strftime('%d %b %Y')
                })
            except Exception as e:
                print(f"Error processing notulensi {notulensi.id}: {e}")
                import traceback
                traceback.print_exc()
                continue
        
        # Always return valid JSON with 200 status
        return jsonify({
            'success': True,
            'upcoming': upcoming_data,
            'recent': recent_data
        }), 200
        
    except Exception as e:
        # Log the full error
        print(f"CRITICAL News feed error: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        
        # Return empty data instead of error to prevent UI breaking
        return jsonify({
            'success': True,  # Still return success to prevent error message
            'upcoming': [],
            'recent': [],
            'error': str(e)
        }), 200  # Use 200 status to prevent JS error handling
    
@app.errorhandler(403)
def forbidden(e):
    return render_template('403.html'), 403

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)